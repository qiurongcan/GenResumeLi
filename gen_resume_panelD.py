# 2025年 生成word版本的简历

# 简历生成分配
import pandas as pd
import os
import random
import subprocess
import glob
from tqdm import tqdm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import json
from box import Box
from pprint import pprint

import random
# 生成第一学历
dir_path = r'Resume_PanelD'
os.makedirs(dir_path, exist_ok=True)

inform = Box.from_json(filename="information.json")
# pprint(inform)

sheet = 'PanelD_第一学历'

df = pd.read_excel(r'简历编码目录.xlsx', index_col=None, sheet_name=sheet)
# print(df.head(3))

# 计算机项目经历


for i in tqdm(range(len(df))):

    resume_id = df.iloc[i,0]
    major = df.iloc[i, 1]
    sex = df.iloc[i, 2]
    degree = "硕士" # 默认都是硕士
    degreeA = df.iloc[i, 3]
    degreeB = df.iloc[i, 4]
    nation = "汉族"

    # 根据性别随机生成名字
    if nation == "汉族":
        name = random.choice(inform.name[sex])
    else:
        name = inform.minority[sex][nation]

    # 出生年月
    birth_year = inform.birth_year[degree] 
    birth_month = random.choice(inform.birth_month)
    
    # 出生地
    birth_place = random.choice(inform.birth_place)

    # 现居地
    living_place = random.choice(inform.birth_place)
    district = random.choice(inform.district[living_place])
    
    
    # 创建新文档
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') # style，所有文字

    style = doc.styles['Heading 1']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') # style，所有文字
    style.font.size = Pt(18)

    h0 = doc.add_heading(f"COSER2025 {sheet} {major}", level=1)
    h0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    h1 = doc.add_heading("1 个人信息", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"""
        头像：统一按性别使用给定头像
        姓名：{name}
        性别：{sex}
        当前身份：应届毕业生
        出生年月：{birth_year}年{birth_month}月
        现居住城市：{living_place}-{district}
        户口所在地：{birth_place}
        政治面貌：共青团员
        手机号码：与注册号码一致
        电子邮箱：与注册邮箱一致（空着不写）
        微信号：与注册微信一致（空着不写）
    """)

    # 个人优势
    personAdvantage = inform.project[major]['personAdvantage']

    h1 = doc.add_heading("2 个人优势", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(personAdvantage)

    h1 = doc.add_heading("3 求职状态", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("在校-正在找工作")

    random_city = random.choice(inform.desire_city)

    h1 = doc.add_heading("4 求职意向", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"""
    期望职位：{inform.tgt_career[major].desire_career}
    期望行业：{inform.tgt_career[major].desire_industry}
    求职偏好：空着不写
    工作城市：北京、上海、广州、西安、{random_city}
    薪资要求：
    工作性质：全职
    """)
    # {inform.desire_salary[major][level+degree]}

    # 工作实习经历
    proj = inform.project[major]
    
    h1 = doc.add_heading("5 工作实习经历", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"""
    职位名称：{proj["jobName"]}
    公司名称：{proj['companyName']}
    所属行业：{proj['industry']}
    在职时间：{proj['workTime']}
    工作内容：
    {proj['workDescription']}
    拥有技能：{proj["skills"]}
    当时月薪：{proj['salary']}
    对这家公司隐藏我的信息（开启）
    """)

    # 项目经历,复用proj
    h1 = doc.add_heading("6 项目经历", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(proj["experience"])

    schoolA = random.choice(inform.schools[living_place][degreeA])
    schoolB = random.choice(inform.schools[living_place][degreeB])

    if degree == "硕士":
        edu_exp = f"""
    （本科阶段）
    学历：本科-统招
    学校名称：{schoolA}
    所学专业：{inform.major[major]}
    在校时间：2019.9-2023.6

    （硕士阶段）
    学历：硕士-统招
    学校名称：{schoolB}
    所学专业：{inform.major[major]}
    在校时间：2023.9-2026.6
"""
    else:
        edu_exp = f"""
    （本科阶段）
    学历：本科-统招
    学校名称：{schoolA}
    所学专业：{major}
    在校时间：2022.9-2026.6        
    """

    h1 = doc.add_heading("7 教育经历", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(edu_exp)

    h1 = doc.add_heading("8 专业技能", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(proj["professionalSkills"])


    h1 = doc.add_heading("9 资格证书", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(inform.certificate)


    h1 = doc.add_heading("10 学生干部经历", level=2)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph = doc.add_paragraph(inform.school_exp)

    # 设置字体
    # for paragraph in doc.paragraphs:
    # for run in paragraph.runs:
    #     # run.font.name = u'宋体'
    #     run.font.size = Pt(14)

    # if major == "STEM":
    #     major_id = 1
    # else:
    #     major_id = 2

    doc.save(f"{dir_path}/{resume_id}.docx")
    # break
    # if i == 10:
    #     break








